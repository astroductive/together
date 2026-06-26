# -*- coding: utf-8 -*-
"""
Builds Together_Graduation_Thesis.docx — the complete FUE-format graduation thesis
for the project "Together: An AI-Based Sign-Language Translator (ASL & ArSL)".

Everything technical is grounded in the repository's actual code. Figures are the
PNGs rendered under thesis/diagrams/. Inline {{GAP|VERIFY|PLACEHOLDER: ...}} markers
and [n] IEEE citations are rendered specially.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
DIAG = os.path.join(HERE, "diagrams")
ASSETS = os.path.join(HERE, "assets")
REPO = os.path.dirname(HERE)
OUT = os.path.join(HERE, "Together_Graduation_Thesis.docx")

NAVY = RGBColor(0x11, 0x23, 0x3A)
BLUE = RGBColor(0x2F, 0x6D, 0xB5)
WARN = RGBColor(0xB0, 0x1A, 0x1A)
GREYTXT = RGBColor(0x44, 0x44, 0x44)

BODY_FONT = "Times New Roman"

# ───────────────────────── low-level helpers ─────────────────────────

def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _set_table_borders(table, color="9BB3CC", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _field(paragraph, instr, placeholder="", run_props=None):
    """Insert a Word field (begin/instr/separate/placeholder/end)."""
    r1 = paragraph.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r1._r.append(fb)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r3._r.append(fs)
    r4 = paragraph.add_run(placeholder)
    if run_props:
        run_props(r4)
    r5 = paragraph.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r5._r.append(fe)


def _seq(paragraph, label, cached="0"):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" SEQ {label} \\* ARABIC ")
    t = OxmlElement("w:t"); t.text = str(cached)
    nr = OxmlElement("w:r"); nr.append(t); fld.append(nr)
    paragraph._p.append(fld)


def _bookmark(paragraph, name, bid):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _cite(paragraph, n):
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), f"ref{n}")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle"); rstyle.set(qn("w:val"), "Hyperlink"); rpr.append(rstyle)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "2F6DB5"); rpr.append(color)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = f"[{n}]"
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


_TOKEN = re.compile(r'(\[\d+\]|\{\{(?:GAP|VERIFY|PLACEHOLDER):[^}]*\}\}|\*\*[^*]+\*\*)')


def _emit_runs(paragraph, text):
    for piece in _TOKEN.split(text):
        if not piece:
            continue
        if re.fullmatch(r'\[\d+\]', piece):
            _cite(paragraph, int(piece[1:-1]))
        elif piece.startswith("{{"):
            inner = piece[2:-2]
            run = paragraph.add_run("⚠ " + inner)
            run.bold = True
            run.font.color.rgb = WARN
        elif piece.startswith("**"):
            run = paragraph.add_run(piece[2:-2]); run.bold = True
        else:
            paragraph.add_run(piece)


class Thesis:
    def __init__(self):
        self.doc = Document()
        self.fig_no = 0
        self.tab_no = 0
        self.figures = []   # (num, caption)
        self.tables = []    # (num, caption)
        self._setup_styles()

    # ---------- styles & page ----------
    def _setup_styles(self):
        d = self.doc
        normal = d.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(12)
        rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
        rpr.set(qn("w:ascii"), BODY_FONT); rpr.set(qn("w:hAnsi"), BODY_FONT)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(6)

        sizes = {"Heading 1": (24, NAVY), "Heading 2": (16, NAVY),
                 "Heading 3": (13, NAVY), "Heading 4": (12, NAVY)}
        for name, (sz, col) in sizes.items():
            st = d.styles[name]
            st.font.name = BODY_FONT
            st.font.size = Pt(sz)
            st.font.bold = True
            st.font.color.rgb = col
            st.paragraph_format.space_before = Pt(14)
            st.paragraph_format.space_after = Pt(8)
            st.paragraph_format.keep_with_next = True
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # page geometry (US Letter, FUE body margins)
        sec = d.sections[0]
        sec.page_width = Inches(8.5); sec.page_height = Inches(11)
        sec.top_margin = Inches(1.25); sec.bottom_margin = Inches(1.25)
        sec.left_margin = Inches(1.25); sec.right_margin = Inches(0.75)

    # ---------- generic content ----------
    def h1(self, text, page_break=True):
        if page_break:
            self.doc.add_page_break()
        return self.doc.add_heading(text, level=1)

    def h2(self, text):
        return self.doc.add_heading(text, level=2)

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    def h4(self, text):
        return self.doc.add_heading(text, level=4)

    def p(self, text, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        para = self.doc.add_paragraph(style=style)
        para.alignment = align
        _emit_runs(para, text)
        return para

    def bullet(self, text):
        para = self.doc.add_paragraph(style="List Bullet")
        _emit_runs(para, text)
        return para

    def numbered(self, text):
        para = self.doc.add_paragraph(style="List Number")
        _emit_runs(para, text)
        return para

    # ---------- figures & tables ----------
    def figure(self, filename, caption, width_in=6.0, max_h=8.2):
        path = os.path.join(DIAG, filename)
        self.fig_no += 1
        with Image.open(path) as im:
            w, h = im.size
        ar = h / w
        wi = width_in
        if wi * ar > max_h:
            wi = max_h / ar
        pc = self.doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.add_run().add_picture(path, width=Inches(wi))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run("Figure ")
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREYTXT
        _seq(cap, "Figure", cached=self.fig_no)
        r2 = cap.add_run(": " + caption)
        r2.font.size = Pt(10); r2.font.color.rgb = GREYTXT
        self.figures.append((self.fig_no, caption))

    def table_caption(self, caption):
        self.tab_no += 1
        cap = self.doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run("Table ")
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREYTXT
        _seq(cap, "Table", cached=self.tab_no)
        r2 = cap.add_run(": " + caption)
        r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = GREYTXT
        self.tables.append((self.tab_no, caption))

    def make_table(self, headers, rows, widths=None, font_size=10, header_bg="2F6DB5"):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        _set_table_borders(t)
        hdr = t.rows[0].cells
        for i, htxt in enumerate(headers):
            hdr[i].text = ""
            para = hdr[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(htxt); run.bold = True; run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(hdr[i], header_bg)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                para = cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                _emit_runs(para, str(val))
                for rn in para.runs:
                    rn.font.size = Pt(font_size)
        if widths:
            for i, wdt in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(wdt)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def save(self):
        self.doc.save(OUT)


# ============================ build the document ============================
T = Thesis()
doc = T.doc

# import the content builders (kept in a second module for readability)
import thesis_content
thesis_content.build(T)

T.save()
print("SAVED", OUT)
print("figures:", T.fig_no, "tables:", T.tab_no)
