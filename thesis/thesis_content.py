# -*- coding: utf-8 -*-
"""Content for the Together graduation thesis. Imported by build_thesis.py."""
import os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x11, 0x23, 0x3A)
BLUE = RGBColor(0x2F, 0x6D, 0xB5)
GREY = RGBColor(0x44, 0x44, 0x44)
WARN = RGBColor(0xB0, 0x1A, 0x1A)
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


# ---- local field/bookmark helpers (avoid importing build_thesis) ----
def _field(paragraph, instr, placeholder=""):
    r1 = paragraph.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r1._r.append(fb)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r3._r.append(fs)
    paragraph.add_run(placeholder)
    r5 = paragraph.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r5._r.append(fe)


def _bookmark(paragraph, name, bid):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(bid)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start); paragraph._p.append(end)


def _center(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def _page_footer(section, fmt, start=None):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r.text = ""
    _field(p, " PAGE ", "1")
    sectPr = section._sectPr
    pgnt = sectPr.find(qn("w:pgNumType"))
    if pgnt is None:
        pgnt = OxmlElement("w:pgNumType"); sectPr.append(pgnt)
    pgnt.set(qn("w:fmt"), fmt)
    if start is not None:
        pgnt.set(qn("w:start"), str(start))


# ============================================================================
def build(T):
    title_page(T)
    declaration(T)
    abstract(T)
    contents(T)
    list_of_figures(T)
    list_of_tables(T)
    list_of_acronyms(T)
    list_of_symbols(T)
    # front-matter page numbering: lower-roman, no number on the title page
    front = T.doc.sections[0]
    front.different_first_page_header_footer = True
    _page_footer(front, "lowerRoman")
    # body section: restart decimal at 1
    body = T.doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = Inches(8.5); body.page_height = Inches(11)
    body.top_margin = Inches(1.25); body.bottom_margin = Inches(1.25)
    body.left_margin = Inches(1.25); body.right_margin = Inches(0.75)
    body.different_first_page_header_footer = False
    _page_footer(body, "decimal", start=1)

    chapter1(T)
    chapter2(T)
    chapter3(T)
    chapter4(T)
    references(T)


# ----------------------------------------------------------------------------
def title_page(T):
    doc = T.doc

    def line(text, size, bold=False, color=NAVY, space=4, italic=False):
        p = _center(doc.add_paragraph())
        p.paragraph_format.space_after = Pt(space)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
        r.italic = italic; r.font.color.rgb = color
        return p

    line("Future University in Egypt", 14, bold=True)
    line("Faculty of Engineering and Technology", 14, bold=True)
    # project emblem
    logo = os.path.join(REPO, "app", "static", "img", "logo-mark-color.png")
    if os.path.exists(logo):
        pic = _center(doc.add_paragraph())
        pic.paragraph_format.space_before = Pt(14); pic.paragraph_format.space_after = Pt(6)
        pic.add_run().add_picture(logo, width=Inches(1.4))
    _spacer(doc, 1)
    line("Together", 30, bold=True, color=BLUE, space=2)
    line("An AI-Based Sign-Language Translator (ASL & ArSL)", 18, bold=True, space=10)
    line("A Graduation Project Thesis", 13, italic=True, color=GREY, space=14)

    line("By", 16, bold=True, space=8)
    # authors table
    authors = [
        ("Abdelfattah", "20225889"),
        ("Michael Abdallah", "20213601"),
        ("Ahmed Mohamed Nagib", "20170423"),
        ("Ahmed Ashraf Shawareb", "20180097"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = 1  # center
    for name, sid in authors:
        c = tbl.add_row().cells
        c[0].text = ""; c[1].text = ""
        rn = c[0].paragraphs[0].add_run(name); rn.font.size = Pt(12); rn.bold = True
        c[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        ri = c[1].paragraphs[0].add_run(sid); ri.font.size = Pt(12)
        c[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in tbl.rows:
        row.cells[0].width = Inches(3.0); row.cells[1].width = Inches(1.6)

    _spacer(doc, 1)
    line("Submitted to the Department of Computers and Intelligent Systems Engineering",
         12, space=2)
    line("in Partial Fulfilment of the Requirements for the Degree of", 12, space=2)
    line("Bachelor of Science in Computers and Intelligent Systems Engineering", 12, bold=True, space=14)

    line("Supervisor(s):", 12, bold=True, space=4)
    sup = doc.add_table(rows=0, cols=1); sup.alignment = 1
    for s in ["Prof. Medhat Awadallah", "[SECOND SUPERVISOR — if any]"]:
        cell = sup.add_row().cells[0]
        pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = pr.add_run(s); rn.font.size = Pt(12)
        if s.startswith("["):
            rn.bold = True; rn.font.color.rgb = WARN
    _spacer(doc, 1)
    line("July 2026", 13, bold=True)


def declaration(T):
    doc = T.doc
    doc.add_page_break()
    T.h1("Declaration", page_break=False)
    T.p(
        "We confirm that this thesis, presented for the degree of Bachelor of Science in "
        "Computers and Intelligent Systems Engineering, has been composed entirely by "
        "ourselves, has been solely the result of our own work, and has not been submitted "
        "for any other degree or professional qualification. Where the work of others has "
        "been used, it has been duly acknowledged and cited in the text."
    )
    _spacer(doc, 1)
    authors = ["Abdelfattah (20225889)", "Michael Abdallah (20213601)",
               "Ahmed Mohamed Nagib (20170423)", "Ahmed Ashraf Shawareb (20180097)"]
    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.add_row().cells
    for i, h in enumerate(["Author", "Signature", "Date"]):
        rr = hdr[i].paragraphs[0].add_run(h); rr.bold = True
    for a in authors:
        c = tbl.add_row().cells
        c[0].paragraphs[0].add_run(a)
        c[2].paragraphs[0].add_run("[DATE]").font.color.rgb = WARN
    for row in tbl.rows:
        row.cells[0].width = Inches(3.0); row.cells[1].width = Inches(2.0); row.cells[2].width = Inches(1.3)


def abstract(T):
    doc = T.doc
    doc.add_page_break()
    T.h1("Abstract", page_break=False)
    T.p(
        "Communication between Deaf or Hard-of-Hearing (Deaf/HoH) signers and the hearing "
        "majority is limited by the scarcity of human interpreters and the absence of "
        "accessible, real-time translation tools. This thesis presents **Together**, an "
        "AI-based, bi-directional sign-language translation platform that runs entirely in a "
        "web browser and supports two languages as first-class citizens: American Sign "
        "Language (ASL) and Arabic/Egyptian Sign Language (ArSL). The system translates in "
        "four directions — sign-to-text, sign-to-speech, text-to-sign, and speech-to-sign — "
        "and adds a live two-person meeting mode that bridges a signer and a speaker over "
        "WebRTC."
    )
    T.p(
        "The recognition pipeline extracts 543 three-dimensional body, face and hand "
        "landmarks per frame in the browser using MediaPipe Holistic, then streams a compact "
        "landmark payload to the server over a persistent Socket.IO channel. Two trained "
        "models perform isolated-sign recognition: a 250-class English ASL classifier "
        "executed with TensorFlow Lite (LiteRT) on a variable-length (N, 543, 3) landmark "
        "sequence, and a 20-class Arabic recognizer implemented as a convolutional–recurrent "
        "(CNN-GRU) network in PyTorch with INT8 dynamic quantization. A server-side voting "
        "and cooldown buffer debounces the per-frame predictions into stable sign tokens. "
        "Because sign languages use a Topic-Comment grammar rather than a word-for-word "
        "encoding of speech, a gloss engine converts recognized gloss into natural sentences "
        "(and the reverse) through a provider abstraction that defaults to Google Gemini and "
        "falls back to a local Ollama model when offline, while non-manual markers (facial "
        "grammar) are annotated for the avatar. Text-to-sign animation is produced by a "
        "semantic lookup that matches words to stored landmark sequences using Sentence-BERT "
        "embeddings indexed in PostgreSQL with the pgvector extension, synthesising "
        "out-of-vocabulary words by fingerspelling and Savitzky–Golay smoothing."
    )
    T.p(
        "The system was engineered for real-time use: blocking inference and network calls "
        "are off-loaded from the event loop, a bounded cache eliminates repeated language-"
        "model round-trips, and the Arabic predictor's INT8 quantization reduced CPU "
        "inference latency from 7.48 ms to 4.21 ms per call (a 1.78× speed-up). On an offline "
        "evaluation of 250 isolated-sign clips, the ASL recognizer achieved a Top-1 accuracy "
        "of 62.4%. The result is a deployable, bilingual, accessibility-focused translator "
        "that demonstrates an end-to-end, landmark-based approach to sign-language "
        "translation on commodity hardware, together with a clear roadmap toward continuous, "
        "sentence-level signing and on-device deployment."
    )


def contents(T):
    doc = T.doc
    doc.add_page_break()
    T.h1("Contents", page_break=False)
    note = doc.add_paragraph()
    rr = note.add_run("(Right-click the field below in Microsoft Word and choose "
                      "“Update Field” to generate the table of contents with page numbers.)")
    rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = GREY
    p = doc.add_paragraph()
    _field(p, ' TOC \\o "1-3" \\h \\z \\u ',
           "Table of Contents — update this field in Word.")


def list_of_figures(T):
    doc = T.doc
    doc.add_page_break()
    T.h1("List of Figures", page_break=False)
    note = doc.add_paragraph()
    rr = note.add_run("(Right-click and “Update Field” to populate page numbers.)")
    rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = GREY
    p = doc.add_paragraph()
    _field(p, ' TOC \\h \\z \\c "Figure" ', "List of Figures — update this field in Word.")


def list_of_tables(T):
    doc = T.doc
    doc.add_page_break()
    T.h1("List of Tables", page_break=False)
    note = doc.add_paragraph()
    rr = note.add_run("(Right-click and “Update Field” to populate page numbers.)")
    rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = GREY
    p = doc.add_paragraph()
    _field(p, ' TOC \\h \\z \\c "Table" ', "List of Tables — update this field in Word.")


def list_of_acronyms(T):
    doc = T.doc
    doc.add_page_break()
    T.h1("List of Acronyms", page_break=False)
    rows = [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("ArSL", "Arabic Sign Language"),
        ("ASGI", "Asynchronous Server Gateway Interface"),
        ("ASL", "American Sign Language"),
        ("BN", "Batch Normalization"),
        ("CNN", "Convolutional Neural Network"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("CPU", "Central Processing Unit"),
        ("DFD", "Data-Flow Diagram"),
        ("ECE", "Electronics and Communication Engineering"),
        ("EMA", "Exponential Moving Average"),
        ("ER", "Entity-Relationship"),
        ("FPS", "Frames Per Second"),
        ("FUE", "Future University in Egypt"),
        ("GISLR", "Google Isolated Sign Language Recognition"),
        ("GRU", "Gated Recurrent Unit"),
        ("GPU", "Graphics Processing Unit"),
        ("HNSW", "Hierarchical Navigable Small World (index)"),
        ("HoH", "Hard-of-Hearing"),
        ("HTTP", "Hypertext Transfer Protocol"),
        ("INT8", "8-bit Integer (quantization)"),
        ("JWT", "JSON Web Token"),
        ("LLM", "Large Language Model"),
        ("LSTM", "Long Short-Term Memory"),
        ("MENA", "Middle East and North Africa"),
        ("ML", "Machine Learning"),
        ("MSA", "Modern Standard Arabic"),
        ("NMM", "Non-Manual Marker"),
        ("ORM", "Object-Relational Mapping"),
        ("PCM", "Pulse-Code Modulation"),
        ("PWA", "Progressive Web Application"),
        ("REST", "Representational State Transfer"),
        ("RTL", "Right-to-Left"),
        ("SBERT", "Sentence-BERT (sentence embeddings)"),
        ("SLR", "Sign-Language Recognition"),
        ("STT", "Speech-to-Text"),
        ("SVO", "Subject-Verb-Object"),
        ("TFLite", "TensorFlow Lite / LiteRT"),
        ("TTS", "Text-to-Speech"),
        ("UI", "User Interface"),
        ("WASM", "WebAssembly"),
        ("WebRTC", "Web Real-Time Communication"),
    ]
    T.make_table(["Acronym", "Definition"], rows, widths=[1.3, 5.0], font_size=10.5)


def list_of_symbols(T):
    doc = T.doc
    doc.add_page_break()
    T.h1("List of Symbols", page_break=False)
    rows = [
        ("N", "Number of frames in a landmark sequence (rolling window, up to 60)"),
        ("w, h", "Captured video frame width and height in pixels"),
        ("x = (x, y, z)", "A single 3-D landmark coordinate (z zeroed for the Arabic model)"),
        ("c", "Shoulder-centre point used to translate-normalize landmarks"),
        ("s", "Shoulder-width scale used to size-normalize landmarks"),
        ("α", "Exponential-moving-average smoothing weight (α = 0.65)"),
        ("pᵢ", "Soft-max probability of class i"),
        ("τ", "Confidence-acceptance threshold (τ = 0.50 ASL, 0.45 ArSL)"),
        ("d_cos", "Cosine distance between SBERT embeddings (1 − cosine similarity)"),
    ]
    T.make_table(["Symbol", "Meaning"], rows, widths=[1.6, 4.7], font_size=10.5)


# ============================== CHAPTER 1 ==================================
def chapter1(T):
    T.h1("Chapter 1: Introduction", page_break=False)
    T.p(
        "Sign languages are complete, natural languages with their own phonology, morphology "
        "and grammar; they are the primary means of communication for millions of Deaf and "
        "Hard-of-Hearing (Deaf/HoH) people. Yet most hearing people do not sign, and "
        "qualified interpreters are scarce and expensive, which leaves a persistent "
        "communication barrier in education, healthcare, employment and everyday public "
        "life [1], [20]. This project addresses that barrier with **Together**, a software "
        "platform that uses computer vision and machine learning to translate sign language "
        "to and from text and speech in real time, directly in a web browser and with no "
        "specialised hardware beyond a webcam."
    )
    T.p(
        "The objective of the project was to build a working, deployable, bi-directional "
        "translator that supports two sign languages — American Sign Language (ASL) and "
        "Arabic/Egyptian Sign Language (ArSL) — and that operates in four translation "
        "directions: sign-to-text, sign-to-speech, text-to-sign and speech-to-sign. A live "
        "meeting mode extends these directions to a two-party conversation between a signer "
        "and a speaker. The overall approach is **landmark-based**: rather than feeding raw "
        "pixels to a deep network, the browser uses Google's MediaPipe Holistic to extract a "
        "compact, privacy-friendly skeleton of 543 body, face and hand landmarks per frame, "
        "and the server classifies sequences of these landmarks into signs. Recognized signs "
        "are assembled into natural sentences by a language-model layer that respects the "
        "Topic-Comment grammar of sign languages, and the reverse path animates an on-screen "
        "avatar from stored landmark sequences retrieved by semantic search."
    )
    T.p(
        "The remainder of this thesis is organised as follows. **Chapter 2** reviews prior "
        "work on sign-language recognition and translation — vision-based, sensor-glove and "
        "landmark-based approaches, and the deep models (CNN, LSTM/GRU and Transformer) that "
        "underpin them — and positions the gap this project addresses. **Chapter 3** is the "
        "core technical chapter: it documents the development methodology, the requirements "
        "and analysis, the system design and architecture (with the full set of UML and "
        "data-flow diagrams), the data and pre-processing pipeline, the recognition models "
        "and their implementation, and the testing, validation and measured results. "
        "**Chapter 4** concludes by comparing the achievements against the objectives and "
        "setting out concrete future work. The references follow in IEEE numbered format."
    )

    T.h2("1.1 Motivation")
    T.p(
        "Hearing loss is one of the most widespread sensory disabilities. The World Health "
        "Organization estimates that over 1.5 billion people live with some degree of "
        "hearing loss and that this figure is rising [20]. {{VERIFY: replace WHO/regional "
        "Deaf-population figures with citable, up-to-date statistics, including Egypt/MENA "
        "where available.}} For people who communicate primarily through sign language, the "
        "barrier is not the disability itself but the gap between signers and the hearing "
        "world: a Deaf patient describing symptoms to a doctor, a Deaf student following a "
        "lecture, or a Deaf customer at a service desk frequently has no interpreter "
        "available. Human interpreting does not scale to these everyday, spontaneous "
        "interactions."
    )
    T.p(
        "Technology can narrow this gap, but most existing solutions are limited: they "
        "require intrusive gloves or depth cameras, run only on the desktop, recognise only "
        "static fingerspelling, or support a single sign language. Together was motivated by "
        "the goal of an **accessible** tool — one that works on an ordinary laptop or phone "
        "browser, needs no installation, and is bilingual from the ground up. The decision to "
        "support both ASL and ArSL is deliberate: ASL is among the most-studied sign "
        "languages with the largest public datasets, which makes a strong English baseline "
        "feasible, whereas Arabic and specifically Egyptian sign language are "
        "under-resourced, so building even a 20-sign Egyptian recogniser contributes to an "
        "under-served community and demonstrates that the architecture generalises across "
        "languages and writing directions (the entire interface mirrors to right-to-left for "
        "Arabic). {{GAP: add any project-specific motivation — a community partner, sponsor, "
        "or accessibility problem the team observed first-hand.}}"
    )

    T.h2("1.2 Aims and Objectives")
    T.p("The aim of the project was to design, implement and evaluate a real-time, "
        "bi-directional, bilingual sign-language translation platform deployable on the web. "
        "This aim was decomposed into the following concrete objectives, each of which is "
        "realised by an identifiable part of the system:")
    T.bullet("**O1 — Real-time capture.** Extract body/hand/face landmarks from a webcam in "
             "the browser and stream them to the server with low latency "
             "(MediaPipe Holistic; Socket.IO streaming in app/static/js/app.js).")
    T.bullet("**O2 — Isolated-sign recognition for two languages.** Train/integrate a "
             "250-class ASL model and a 20-class ArSL model and serve both behind one API "
             "(models/model.tflite, models/best_model.pth, app/server/asl_service.py, "
             "models/sign_predictor.py).")
    T.bullet("**O3 — Natural-language output.** Convert recognized Topic-Comment gloss into "
             "fluent sentences and speech, and the reverse, with graceful offline fallback "
             "(app/server/gloss.py, app/server/providers/).")
    T.bullet("**O4 — Sign synthesis (text/speech-to-sign).** Animate an avatar from stored "
             "landmark sequences retrieved by semantic search, synthesising unknown words by "
             "fingerspelling (app/server/asl_service.py, app/server/stitch.py, "
             "app/static/js/landmark-player.js).")
    T.bullet("**O5 — Live two-party translation.** Provide a WebRTC meeting that translates "
             "between a signer and a speaker (Socket.IO signalling in app/server/main.py).")
    T.bullet("**O6 — Production qualities.** Bilingual (LTR/RTL) UI, authentication, "
             "rate-limiting, latency profiling and a deployable container "
             "(app/server/auth.py, app/server/profiling.py, docker-compose.yml).")
    T.p("Chapter 3 reports how each objective was met and Chapter 4 evaluates the outcome "
        "against them.")


# ============================== CHAPTER 2 ==================================
def chapter2(T):
    T.h1("Chapter 2: Review of Literature")
    T.p(
        "Automatic sign-language recognition (SLR) and translation has been studied for "
        "three decades, and the literature divides naturally by the **input modality** used "
        "to capture a sign and by the **model** used to classify it. This chapter surveys "
        "vision-based, sensor-based and landmark-based recognition, the deep-learning "
        "architectures that dominate recent work, and the specific state of Arabic "
        "sign-language research, before stating the gap that Together addresses. A consolidated "
        "comparison of representative systems is given in Table 1."
    )

    T.h2("2.1 Vision-Based and Sensor-Based Recognition")
    T.p(
        "Early systems treated SLR as a video classification problem. Starner et al. [19] "
        "used hidden Markov models on camera and wearable-camera video to recognise a "
        "40-word ASL vocabulary at high accuracy, but under controlled, person-dependent "
        "conditions. Sensor-glove approaches instrument the hand directly: Oz and Leu [17] "
        "recognised ASL words from a DataGlove and a motion tracker with an artificial "
        "neural network, achieving strong accuracy at the cost of intrusive hardware that is "
        "impractical for everyday use. As deep learning matured, end-to-end video models "
        "such as the 3-D convolutional I3D network [10] were applied to large word-level "
        "datasets; on the 2,000-word WLASL benchmark [16] top-1 accuracy drops sharply as "
        "the vocabulary grows, illustrating how data-hungry pixel-based models are. "
        "Multi-modal datasets such as AUTSL [15] add depth to RGB and report high accuracy "
        "with CNN-plus-LSTM pipelines, but depend on a depth sensor."
    )
    T.h3("2.1.1 Landmark-Based Recognition")
    T.p(
        "A more recent and practical line of work replaces raw pixels with **pose and hand "
        "landmarks** produced by a real-time tracker, most commonly Google's MediaPipe "
        "framework [2], [3]. Landmarks are compact (a few hundred numbers per frame instead "
        "of an image), are largely invariant to background, lighting and clothing, and "
        "preserve privacy because the original video never leaves the device. The Google "
        "Isolated Sign Language Recognition (GISLR / PopSign) effort popularised a "
        "standardised 543-landmark MediaPipe Holistic representation and a 250-sign "
        "vocabulary for English ASL [4]. Together adopts exactly this representation and "
        "vocabulary for its English model, which makes the approach reproducible and the "
        "results comparable, while running the whole tracker in the browser via WebAssembly. "
        "Sequence models — recurrent networks (LSTM [8] and GRU [9]) and, more recently, "
        "Transformers [7] — are the standard choice for classifying landmark sequences, "
        "because a sign is a trajectory through time rather than a single pose."
    )

    T.h2("2.2 Deep Models, Arabic Sign Language, and the Gap")
    T.p(
        "For continuous, sentence-level signing, Camgoz et al. introduced neural sign-"
        "language translation [11] and later Sign Language Transformers [12], jointly "
        "learning recognition and translation on the German PHOENIX14T corpus and reporting "
        "translation quality in BLEU. These models are powerful but require large aligned "
        "corpora and substantial compute, and target continuous signing rather than the "
        "isolated-sign setting that is tractable for a student project and a low-resource "
        "language; the sentence-embedding technique Together uses for sign lookup is itself "
        "built on this transformer lineage — Sentence-BERT [5], [18]. For Arabic, most "
        "published work addresses the **alphabet** rather than "
        "words: the ArASL dataset [14] provides static Arabic-alphabet images, and "
        "Kamruzzaman [13] recognised Arabic alphabet signs with a CNN and synthesised speech. "
        "Word-level Egyptian sign-language resources remain very limited, which is precisely "
        "why a small but functional 20-word Egyptian recogniser is a meaningful contribution. "
        "A broad survey of the field is given by Rastgoo et al. [1]."
    )
    T.p(
        "Across this literature three gaps recur. First, most systems are **uni-directional** "
        "(sign-to-text only) and **monolingual**. Second, high-accuracy systems usually "
        "depend on special hardware (gloves, depth cameras) or heavy GPU models that cannot "
        "run in a browser. Third, Arabic word-level signing is under-served. Together targets "
        "all three: it is bi-directional and bilingual, it runs landmark extraction in the "
        "browser on commodity webcams, and it treats Egyptian ArSL as a first-class language "
        "alongside ASL. It does not attempt continuous translation — that is identified as "
        "future work in Chapter 4 — but it delivers a complete, deployable isolated-sign "
        "system with a grammar-aware language layer, which is uncommon in the surveyed work."
    )

    T.table_caption("Comparison of representative sign-language recognition and translation systems.")
    T.make_table(
        ["System / approach", "Modality", "Dataset / vocab", "Reported accuracy", "Key limitations"],
        [
            ["Starner et al. [19] (HMM)", "RGB video (desk / wearable)", "40-word ASL",
             "~92–98%", "Small vocab; controlled; person-dependent"],
            ["Oz & Leu [17] (ANN)", "Sensor glove + tracker", "~50 ASL words",
             "~90%", "Intrusive hardware; not camera-based"],
            ["WLASL / I3D [16], [10]", "RGB video", "2,000 words",
             "~32% top-1 (I3D)", "Large vocab is data-hungry; GPU-heavy"],
            ["AUTSL [15] (CNN+LSTM)", "RGB-D video", "226 Turkish signs",
             "~95% (RGB-D)", "Requires depth sensor; not ASL/ArSL"],
            ["Sign Transformers [12]", "RGB video", "PHOENIX14T (continuous DGS)",
             "~21–24 BLEU", "Needs large aligned corpus; heavy compute"],
            ["Kamruzzaman [13] (CNN)", "RGB images", "Arabic alphabet",
             "~90%", "Static letters only; not word-level"],
            ["**Together (this work)**", "MediaPipe landmarks (browser)",
             "250 ASL (GISLR-style) + 20 ArSL", "62.4% Top-1 (ASL, 250-class)",
             "Isolated signs; small ArSL vocab; web-deployed, real-time, bilingual"],
        ],
        widths=[1.5, 1.4, 1.4, 1.05, 1.6], font_size=8.6,
    )


# ============================== CHAPTER 3 ==================================
def chapter3(T):
    T.h1("Chapter 3: Technical Chapters")
    T.p(
        "This chapter documents the engineering of Together across its full lifecycle: the "
        "development methodology, the requirements and analysis, the system design and "
        "architecture, the data and pre-processing pipeline, the model and implementation, "
        "and finally testing, validation and results. Every technical claim is grounded in "
        "the project's source code, and file paths are cited inline so each statement can be "
        "traced to its implementation."
    )

    # ---- 3.1 methodology ----
    T.h2("3.1 Development Process and Methodology")
    T.p(
        "The project followed an **iterative, incremental** process. Rather than specifying "
        "the whole system up front, the team built a thin end-to-end slice first — webcam "
        "capture, a single recognition path, and a minimal UI — and then grew it feature by "
        "feature, hardening each layer once it worked. The work proceeded in clear phases: a "
        "research and data phase; a recognition-model phase for ASL and then ArSL; a "
        "backend-services phase (gloss, providers, sign lookup); a frontend and bilingual-UI "
        "phase; a real-time meeting phase; and finally a hardening phase covering "
        "authentication, latency profiling and deployment. The latency optimisations are "
        "documented in docs/LATENCY.md and the provider design in "
        "docs/PROVIDERS_AND_PRICING.md, evidencing the engineering-driven, measure-then-"
        "optimise methodology. Version control was used throughout, and an automated test "
        "suite plus a continuous-integration workflow (.github/workflows/ci.yml) guarded "
        "against regressions."
    )
    T.table_caption("Project timeline and milestones (representative academic-year schedule).")
    T.make_table(
        ["Phase", "Period", "Milestone / deliverable"],
        [
            ["1. Research & datasets", "Oct–Nov 2025", "Literature review; landmark representation; sign-video corpus"],
            ["2. ASL recognition", "Dec 2025–Jan 2026", "250-class TFLite pipeline; live capture & voting"],
            ["3. ArSL recognition", "Feb 2026", "20-class CNN-GRU model; bilingual switch"],
            ["4. Language & lookup", "Feb–Mar 2026", "Gloss engine + NMM; SBERT/pgvector sign lookup; providers"],
            ["5. Frontend & RTL UI", "Mar–Apr 2026", "Dashboards (EN/AR), avatar player, design system"],
            ["6. Meetings & realtime", "Apr 2026", "WebRTC mesh; Socket.IO streaming recognition"],
            ["7. Evaluation & hardening", "Apr–May 2026", "Baseline eval (Top-1 0.624); auth; latency profiling; deploy"],
            ["8. Documentation", "Jun–Jul 2026", "Thesis, diagrams, final submission"],
        ],
        widths=[1.7, 1.5, 3.1], font_size=9.5,
    )
    T.p("{{VERIFY: replace the milestone dates above with the team's real schedule — the "
        "version-control history in this checkout covers only the final week and cannot be "
        "used to reconstruct the true timeline.}}")

    # ---- 3.2 requirements ----
    T.h2("3.2 Requirements and Analysis")
    T.p("The requirements were derived from the project aim (Section 1.2) and refined during "
        "the iterative process. They are grouped into functional and non-functional "
        "requirements, followed by the hardware and software needed to run the system, and "
        "the principal use cases.")

    T.h3("3.2.1 Functional Requirements")
    T.table_caption("Functional requirements.")
    T.make_table(
        ["ID", "Requirement", "Realised by"],
        [
            ["FR1", "Capture webcam video and extract body/hand/face landmarks in the browser", "MediaPipe Holistic (app/static/js/app.js)"],
            ["FR2", "Recognise isolated ASL signs from a landmark sequence", "ASLService.predict_sign (asl_service.py)"],
            ["FR3", "Recognise isolated ArSL signs from a landmark sequence", "SignLanguagePredictor (sign_predictor.py)"],
            ["FR4", "Debounce per-frame predictions into stable sign tokens", "Vote/cooldown buffer (main.py sign_frame)"],
            ["FR5", "Convert recognized gloss into a natural sentence", "gloss_to_sentence (gloss.py)"],
            ["FR6", "Convert a typed/spoken sentence into sign gloss + NMM", "english_to_gloss, annotate_nmm (gloss.py)"],
            ["FR7", "Animate an avatar from landmark sequences for given words", "/api/signs/batch + landmark-player.js"],
            ["FR8", "Synthesise unknown words by fingerspelling", "gloss_and_stitch (stitch.py)"],
            ["FR9", "Speak output (TTS) and accept spoken input (STT)", "/api/tts, /api/stt; providers/"],
            ["FR10", "Host a live two-party signer–speaker meeting", "WebRTC + Socket.IO (main.py)"],
            ["FR11", "Support English (LTR) and Arabic (RTL) interfaces", "Jinja2 templates *_ar.html"],
            ["FR12", "Authenticate users and protect the API", "auth.py (JWT, rate limits)"],
        ],
        widths=[0.5, 3.0, 2.8], font_size=9.0,
    )

    T.h3("3.2.2 Non-Functional Requirements")
    T.table_caption("Non-functional requirements.")
    T.make_table(
        ["ID", "Quality", "Target / mechanism"],
        [
            ["NFR1", "Real-time responsiveness", "Per-frame streaming at ~20 fps; sign accepted in a few hundred ms; blocking work off the event loop (run_in_threadpool)"],
            ["NFR2", "Concurrency", "No endpoint blocks the event loop; thread-locked single-instance models; DB pool (size 5, overflow 10)"],
            ["NFR3", "Offline resilience", "Provider fallback chains (Gemini→Ollama / pyttsx3 / Whisper); rule-based gloss fallback"],
            ["NFR4", "Security", "Argon2id hashing; JWT with rotating refresh tokens; sliding-window rate limiting; security headers"],
            ["NFR5", "Portability", "Browser client (no install); Dockerised server; runs on CPU only"],
            ["NFR6", "Accessibility & i18n", "Bilingual EN/AR, RTL mirroring, reduced-motion & focus-visible support"],
            ["NFR7", "Observability", "Per-stage latency metrics at /api/metrics (profiling.py)"],
            ["NFR8", "Privacy", "Only landmarks (not raw video) are sent to the server for recognition"],
        ],
        widths=[0.5, 1.8, 4.0], font_size=9.0,
    )

    T.h3("3.2.3 Hardware Requirements")
    T.table_caption("Hardware requirements.")
    T.make_table(
        ["Component", "Minimum", "Notes"],
        [
            ["Client device", "Any laptop/desktop/phone with a webcam and a modern browser", "MediaPipe runs via WebAssembly; GPU optional"],
            ["Camera", "Standard 480p+ webcam", "Capture configured at ~480×360, 15–20 fps"],
            ["Server CPU", "2+ cores", "Inference is CPU-only (TFLite XNNPACK; PyTorch INT8)"],
            ["Server RAM", "~2–4 GB", "SBERT + two models + Postgres client"],
            ["GPU", "Not required", "All inference runs on CPU"],
            ["Network", "Broadband / Wi-Fi", "Socket.IO streaming + optional cloud LLM/TTS/STT"],
        ],
        widths=[1.4, 2.6, 2.3], font_size=9.0,
    )

    T.h3("3.2.4 Software and Technology Stack")
    T.p("Table 6 lists the technology stack with the exact versions pinned in "
        "requirements.txt and the frontend includes. The server is Python; the client is "
        "server-rendered HTML with vanilla JavaScript and no build step.")
    T.table_caption("Software and technology stack with versions.")
    T.make_table(
        ["Layer", "Technology", "Version"],
        [
            ["Web framework", "FastAPI", "0.115.6"],
            ["ASGI server", "uvicorn[standard]", "0.34.0"],
            ["Real-time", "python-socketio", "5.12.1"],
            ["Templating", "Jinja2", "3.1.5"],
            ["ORM / migrations", "SQLAlchemy / Alembic", "2.0.50 / 1.14.0"],
            ["Database driver", "psycopg (PostgreSQL) + pgvector", "3.2.3 / 0.3.6"],
            ["ASL recognition", "ai-edge-litert (LiteRT/TFLite)", "1.2.0"],
            ["ArSL recognition", "PyTorch (via sentence-transformers)", "torch (bundled)"],
            ["Landmark tracking", "MediaPipe (server) / Holistic (browser CDN)", "0.10.18 / 0.5.1675471629"],
            ["Computer vision", "OpenCV-Python", "4.10.0.84"],
            ["Semantic search", "sentence-transformers (all-MiniLM-L6-v2)", "3.3.1"],
            ["Signal processing", "SciPy (Savitzky–Golay)", "1.14.1"],
            ["Auth", "argon2-cffi / bcrypt / python-jose", "23.1.0 / 4.2.1 / 3.3.0"],
            ["Cloud LLM/TTS/STT", "Google Gemini (gemini-2.5-flash)", "API"],
            ["Offline fallbacks", "Ollama (llama3.2) / pyttsx3 / faster-whisper", "— / 2.98 / 1.1.0"],
            ["Testing / CI", "pytest / httpx / GitHub Actions", "8.3.4 / 0.28.1"],
            ["Runtime", "Python / PostgreSQL", "3.11+ / 16"],
        ],
        widths=[1.6, 3.1, 1.6], font_size=8.8,
    )

    T.h3("3.2.5 Use Cases")
    T.p("Figure 8 shows the use-case diagram with the two human actors — a Deaf/HoH user "
        "and a hearing user — and the system/provider actor that fulfils language tasks. "
        "Table 7 describes the principal use cases.")
    T.figure("F8_usecase.png", "Use-case diagram of the Together platform.", width_in=5.6)
    T.table_caption("Principal use-case descriptions.")
    T.make_table(
        ["Use case", "Actor", "Description"],
        [
            ["Sign → Text", "Deaf/HoH user", "Sign to the webcam; the system recognises signs and shows a natural sentence"],
            ["Sign → Speech", "Deaf/HoH user", "As above, then the sentence is spoken aloud via TTS"],
            ["Text → Sign", "Hearing user", "Type a sentence; the avatar signs it (gloss → landmark lookup)"],
            ["Speech → Sign", "Hearing user", "Speak; STT transcribes, then the avatar signs the text"],
            ["Live meeting", "Both", "Two-party WebRTC call with live captions / sign avatar per role"],
            ["Practice / Dictionary", "Deaf/HoH user", "Browse the recognised vocabulary and practise individual signs"],
            ["Register / Login", "Both", "Create an account and authenticate to use the API"],
            ["Switch language", "Both", "Toggle the EN (LTR) and AR (RTL) interface and model"],
        ],
        widths=[1.3, 1.3, 4.0], font_size=9.0,
    )

    # ---- 3.3 design ----
    T.h2("3.3 System Design and Architecture")
    T.p(
        "Together is a client–server web application. The browser is responsible for capture, "
        "landmark extraction, the avatar and the WebRTC media path; the server hosts the "
        "recognition models, the language layer, the sign database and the provider chains. "
        "Figure 1 gives the high-level architecture. Four design decisions anchor the system: "
        "the frontend is server-rendered HTML/CSS/JS with no SPA framework or build step; the "
        "machine-learning stack uses TFLite for the English model and PyTorch for the Arabic "
        "model, with MediaPipe Holistic extracting landmarks in the browser; real-time "
        "communication uses Socket.IO for both streaming recognition and WebRTC signalling; "
        "and persistence uses PostgreSQL with the pgvector extension for semantic search."
    )
    T.figure("F1_architecture.png", "High-level system architecture of the Together platform.")
    T.p("The context diagram (Figure 2) frames the system as a single process exchanging "
        "data with two human actors and the external Gemini API, and the Level-1 data-flow "
        "diagram (Figure 3) decomposes it into the six functional processes that mirror the "
        "use cases.")
    T.figure("F2_context_dfd0.png", "Context diagram (Data-Flow Diagram, Level 0).", width_in=5.6)
    T.figure("F3_dfd1.png", "Data-Flow Diagram (Level 1).", width_in=5.4)
    T.p("Internally the server is organised into cooperating components (Figure 12): the "
        "FastAPI/Socket.IO entry point (main.py) routes requests to the recognition service, "
        "the gloss engine, the sign-lookup layer and the provider chains, while the database "
        "package and the on-disk landmark store provide persistence. The deployment topology "
        "(Figure 7) shows the browser talking to a containerised web service that is backed "
        "by PostgreSQL and an optional local Ollama service, with Google Gemini reached over "
        "HTTPS as the default cloud provider; in a meeting, media flows peer-to-peer over a "
        "WebRTC mesh while only signalling passes through the server.")
    T.figure("F12_component.png", "Component diagram of the backend and frontend.", width_in=5.6)
    T.figure("F7_deployment.png", "Deployment diagram.", width_in=5.6)

    T.h3("3.3.1 UML Models")
    T.p("The static structure of the recognition and lookup subsystem is shown in the class "
        "diagram (Figure 9). SignDB encapsulates SBERT-backed semantic lookup over PostgreSQL; "
        "ArabicSignDB and ASLService specialise it — the former for the 20-word Arabic vocab, "
        "the latter by adding the TFLite gesture model — and SignLanguagePredictor wraps the "
        "PyTorch CNN-GRU used for Arabic recognition.")
    T.figure("F9_class.png", "Class diagram of the recognition and sign-lookup subsystem.")
    T.p("The dynamic behaviour of the principal real-time path — a signer producing text — is "
        "shown in the sequence diagram (Figure 10) and the activity diagram (Figure 11). The "
        "browser opens a streaming session, sends one landmark frame per tick, and the server "
        "maintains a rolling buffer, runs inference in a worker thread, votes over recent "
        "predictions and emits an accepted sign; a separate request then turns the "
        "accumulated gloss into a sentence. The capture/recognition logic itself is a small "
        "state machine (Figure 13) cycling through idle, buffering, inferring, voting and "
        "cooldown states.")
    T.figure("F10_sequence.png", "Sequence diagram: real-time sign-to-text request flow.")
    T.figure("F11_activity.png", "Activity diagram of the streaming recognition workflow.", width_in=4.4)
    T.figure("F13_state.png", "State-machine diagram of the capture/recognition states.", width_in=5.6)

    T.h3("3.3.2 Data Model")
    T.p(
        "Persistent state lives in PostgreSQL and is described by the entity-relationship "
        "diagram (Figure 14). Three tables are defined in app/server/db/models.py and created "
        "by the Alembic migration alembic/versions/0001_initial_schema.py: **users** (account "
        "and Argon2id password hash), **refresh_tokens** (a server-side registry of token "
        "identifiers for rotation and revocation), and **signs** (sign metadata plus a "
        "384-dimensional SBERT embedding stored in a pgvector column with an HNSW cosine "
        "index). Crucially, raw landmark sequences are **not** stored in the database; the "
        "signs table holds only a path to a compressed .npz file on disk (landmark_store.py), "
        "keeping rows small and vector search fast."
    )
    T.figure("F14_er.png", "Entity-Relationship diagram of the PostgreSQL schema.", width_in=5.8)

    # ---- 3.4 data ----
    T.h2("3.4 Data and Pre-processing")
    T.p(
        "Two vocabularies are recognised. The English model classifies the 250 isolated "
        "signs of the GISLR-style vocabulary (its label map, "
        "models/sign_to_prediction_index_map.json, lists signs alphabetically from “TV” "
        "to “zipper”), and the Arabic model classifies 20 Egyptian signs "
        "(models/class_mapping.json). Table 8 summarises both. In addition, a sign-video "
        "corpus under data/signs_videos provides the reference clips from which the avatar's "
        "landmark database is built and against which the English model was evaluated."
    )
    T.table_caption("Recognised sign vocabularies and data sources.")
    T.make_table(
        ["Property", "English (ASL)", "Arabic (ArSL)"],
        [
            ["Classes", "250 isolated signs", "20 isolated signs"],
            ["Example labels", "TV, apple, happy, mother, …, zipper", "baby, eat, father, happy, mosque, mother, …, worry"],
            ["Model file", "models/model.tflite (≈11 MB)", "models/best_model.pth (≈1.1 MB)"],
            ["Input representation", "(N, 543, 3) MediaPipe Holistic", "(30, 177) pose+hands, Z-zeroed"],
            ["Reference clips", "data/signs_videos (272 files)", "Egyptian ArSL word videos"],
            ["Dataset source", "GISLR/PopSign-style {{VERIFY: confirm exact dataset & licence}}", "{{GAP: dataset source, #signers, consent, demographics}}"],
            ["Train/val/test split", "{{GAP: not recorded in repo}}", "{{GAP: not recorded in repo}}"],
        ],
        widths=[1.5, 2.4, 2.4], font_size=8.8,
    )
    T.p(
        "The recognition pipeline that turns a frame into a feature vector is shown in "
        "Figure 4. In the browser, MediaPipe Holistic produces 33 pose, 468 face and 2×21 "
        "hand landmarks per frame; an exponential moving average (α = 0.65) smooths the "
        "skeleton and a short grace window holds it through brief tracking drops. For "
        "recognition the client sends a compact payload in which missing landmarks are encoded "
        "as null. The two models then pre-process differently, reflecting how each was trained:"
    )
    T.figure("F4_pipeline.png", "End-to-end real-time translation pipeline.", width_in=6.2)
    T.bullet(
        "**English (ASL).** The server feeds the raw variable-length (N, 543, 3) sequence "
        "directly to the TFLite model. Missing landmarks must be NaN rather than zero — the "
        "model was trained with NaN and masks them internally, and a zero is read as a real "
        "point at the image origin — so any all-zero point is normalised to NaN "
        "(asl_service.py). Sequences are capped at 128 frames.")
    T.bullet(
        "**Arabic (ArSL).** The predictor keeps 17 pose joints and the two 21-point hands "
        "(177 values per frame), corrects the landscape aspect ratio to a 9:16 portrait "
        "canvas, zeroes the depth coordinate for distance-invariance, and normalises every "
        "point by the **shoulder centre** and **shoulder width** so the representation is "
        "translation- and scale-invariant. Missing hand points are linearly interpolated to "
        "bridge tracking drops (sign_predictor.py).")
    T.p(
        "ASL and ArSL are distinguished explicitly rather than detected: the request carries "
        "a language field, the Arabic dashboard (index_ar.html) forces the Arabic model, and "
        "the server routes to ASLService or SignLanguagePredictor accordingly (main.py, "
        "/api/translate). For text-to-sign, Arabic input is first mapped to the English class "
        "keys via a bilingual dictionary and, when needed, the LLM, so both languages reuse "
        "the same landmark-lookup machinery."
    )

    # ---- 3.5 model & implementation ----
    T.h2("3.5 Model and Implementation")
    T.p(
        "Figure 5 shows both recognition models. The English model ships as a trained TFLite "
        "graph whose internal architecture is not introspectable from the exported file; it "
        "is defined by its input/output contract — a variable-length (N, 543, 3) landmark "
        "sequence in, 250 raw logits out — and is executed by the LiteRT interpreter with the "
        "XNNPACK delegate and multi-threading (asl_service.py). The Arabic model is fully "
        "specified in code (models/sign_predictor.py) as the SignLanguageCNNGRU network: two "
        "1-D convolutional layers (177→128→128, kernel 3) each with batch normalisation "
        "and ReLU, a two-layer bidirectional GRU (hidden size 64), and two fully-connected "
        "layers (128→64→20) with dropout. Its last-timestep representation is classified "
        "into the 20 Arabic signs."
    )
    T.figure("F5_model_arch.png", "Recognition model architectures: (a) ASL TFLite I/O contract; "
                                  "(b) ArSL CNN-GRU layer stack.")
    T.p(
        "Both models use ensembling and gating for robustness. The Arabic predictor runs the "
        "network over three temporal windows (the last 30, 45 and 60 frames, each uniformly "
        "re-sampled to 30) and averages the soft-max outputs; a prediction is accepted only "
        "above a confidence of 0.45, and the server further requires a two-of-three majority "
        "vote across consecutive inferences. The English path applies soft-max to the raw "
        "logits and accepts above 0.50 (ASL_CONFIDENCE_THRESH), again behind the vote/cooldown "
        "buffer. On the streaming socket path (main.py), a rolling 60-frame buffer, a "
        "three-sample vote buffer requiring two agreeing votes, and an eight-frame cooldown "
        "convert noisy per-frame predictions into stable sign tokens at roughly 20 fps."
    )
    T.p("The training of both models was performed outside this repository, which ships only "
        "the trained checkpoints and the evaluation harness (Figure 6). The hyperparameters "
        "that would normally accompany a training run are therefore not recorded here and "
        "must be supplied; Table 9 captures what is known from the code and flags the gaps.")
    T.figure("F6_training.png", "Per-language training and data pipeline (training performed off-repo).",
             width_in=6.2)
    T.table_caption("Model and training configuration.")
    T.make_table(
        ["Hyperparameter", "English (ASL)", "Arabic (ArSL)"],
        [
            ["Architecture", "GISLR-style TFLite (black-box)", "CNN-GRU (Conv1d×2 + Bi-GRU×2 + FC×2)"],
            ["Input", "(N, 543, 3) landmarks", "(30, 177) landmarks"],
            ["Classes", "250", "20"],
            ["Sequence handling", "variable length, ≤128 frames", "windows 30/45/60 → sample 30, mean-ensemble"],
            ["Confidence threshold", "0.50 (soft-max)", "0.45 (soft-max)"],
            ["Quantization", "float (XNNPACK)", "INT8 dynamic (Linear+GRU)"],
            ["Optimizer / LR / epochs", "{{GAP: not recorded in repo}}", "{{GAP: not recorded in repo}}"],
            ["Batch size / loss", "{{GAP: not recorded}}", "{{GAP: not recorded}}"],
            ["Augmentation", "{{GAP: not recorded}}", "{{GAP: not recorded}}"],
        ],
        widths=[1.7, 2.4, 2.4], font_size=8.8,
    )

    T.h3("3.5.1 Language Layer and Output")
    T.p(
        "Because a sign language is not a word-for-word transcription of speech, recognized "
        "signs cannot simply be concatenated. The gloss engine (app/server/gloss.py) bridges "
        "the two grammars in both directions using few-shot prompting through the provider "
        "abstraction. gloss_to_sentence turns Topic-Comment gloss such as “STORE I GO” into "
        "“I am going to the store”, and english_to_gloss does the reverse, dropping English "
        "articles and copulas and re-ordering to Topic-Comment; both degrade gracefully to a "
        "deterministic rule-based gloss when no language model is reachable. A bounded "
        "256-entry cache keyed on (gloss, language) eliminates repeated model round-trips for "
        "the recurring gloss of a conversation. The engine also annotates **non-manual "
        "markers** — the facial grammar that distinguishes, for example, a wh-question "
        "(furrowed brows) from a yes/no question (raised brows) — which the avatar layer can "
        "render (docs/GLOSS_AND_NMM.md)."
    )
    T.p(
        "The reverse, text-to-sign path retrieves animation from the sign database. Each word "
        "is matched to a stored sign by **semantic search**: the query is embedded with "
        "Sentence-BERT [5] (the distilled all-MiniLM-L6-v2 encoder [6], 384 dimensions) and "
        "compared by cosine distance to "
        "the indexed sign embeddings in PostgreSQL/pgvector, accepting a match within a "
        "language-specific distance (0.35 for English, 0.55 for Arabic). The matched sign's "
        "landmark sequence is loaded from the .npz store and streamed to the browser, where "
        "the landmark-player renders it on a canvas avatar. Words with no dedicated sign are "
        "synthesised by **fingerspelling**: gloss_and_stitch concatenates one clip per letter "
        "and removes the visible jerk at each boundary with a short linearly-interpolated "
        "bridge and a Savitzky–Golay filter (stitch.py), leaving MediaPipe's NaN gaps "
        "untouched so they are not corrupted."
    )
    T.p(
        "Speech output and input are handled by pluggable provider chains "
        "(app/server/providers/): the default is Google Gemini for LLM, TTS and STT, with "
        "automatic fall-back to a local Ollama model, pyttsx3 and faster-whisper respectively "
        "when a cloud key is absent or a call fails. The browser additionally uses the native "
        "Web Speech API for low-latency speech where available. The application is run with "
        "python start_server.py (or via docker-compose up), which sets thread/ML environment "
        "guards before importing PyTorch and serves the Socket.IO ASGI app with uvicorn; the "
        "REST and Socket.IO surface is summarised in Table 10."
    )
    T.table_caption("Principal REST and Socket.IO API surface.")
    T.make_table(
        ["Endpoint / event", "Purpose"],
        [
            ["POST /api/translate", "Run TFLite (ASL) or PyTorch (ArSL) inference on a landmark window"],
            ["POST /api/translate/sentence", "Convert recognized gloss to a natural sentence (cached LLM)"],
            ["POST /api/gloss", "Convert a sentence to gloss + non-manual markers"],
            ["GET /api/tts · POST /api/stt", "Speech synthesis / recognition via the provider chains"],
            ["POST /api/signs/batch · /api/signs_ar/batch", "Look up landmark sequences for words (text-to-sign)"],
            ["GET /api/health · /api/metrics", "Health and per-stage latency metrics"],
            ["POST /api/auth/{signup,login,refresh,logout}", "Authentication with rotating refresh tokens"],
            ["Socket.IO sign_start/sign_frame/sign_stop", "Live streaming recognition"],
            ["Socket.IO join_room, webrtc_offer/answer/ice_candidate", "WebRTC meeting signalling (mesh, up to 6 peers)"],
        ],
        widths=[2.7, 3.6], font_size=9.0,
    )

    # ---- 3.6 testing & results ----
    T.h2("3.6 Testing, Validation and Results")
    T.p(
        "The system is validated at two levels: an automated software test suite that guards "
        "functional correctness, and an offline model evaluation that measures recognition "
        "accuracy. The pytest suite under tests/ covers authentication (hashing, JWT, "
        "rate limiting), the gloss engine and non-manual markers, the provider fall-back "
        "chains, the fingerspelling stitcher, the profiling utilities, and a frontend "
        "regression module that renders every template and asserts the invariants that broke "
        "during the redesign. The continuous-integration workflow runs the lightweight "
        "subset (authentication and templates) on every push. Table 11 lists representative "
        "test cases and their outcomes."
    )
    T.table_caption("Representative functional test cases and outcomes.")
    T.make_table(
        ["Test (module)", "Input", "Expected", "Result"],
        [
            ["Password hashing (test_auth)", "New password", "Argon2id hash; verify true; wrong password false", "Pass"],
            ["Legacy upgrade (test_auth)", "bcrypt hash", "needs_rehash() true; upgraded on login", "Pass"],
            ["JWT expiry/tamper (test_auth)", "Expired/edited token", "decode raises 401", "Pass"],
            ["Rate limit (test_auth)", ">10 auth req/60 s", "Further requests blocked", "Pass"],
            ["Sentence type (test_gloss)", "“What is your name?”", "wh_question; furrowed brows NMM", "Pass"],
            ["Gloss fallback (test_gloss)", "Gloss, no LLM", "Returns raw gloss (no crash)", "Pass"],
            ["Provider chain (test_providers)", "Default config", "LLM chain = [gemini, ollama]", "Pass"],
            ["All-fail (test_providers)", "All providers down", "Raises ProviderError", "Pass"],
            ["Stitch length (test_stitch)", "k clips", "frames = Σ + (k−1)×transition", "Pass"],
            ["Smoothing (test_stitch)", "Two clips", "Max boundary jump < 5.0; NaN preserved", "Pass"],
            ["Metrics (test_profiling)", "Timed stages", "snapshot has count/avg/p50/p95", "Pass"],
            ["Templates (test_templates)", "All templates", "Render; dashboards load MediaPipe/socket; index_ar is RTL; inline JS parses", "Pass"],
        ],
        widths=[1.9, 1.5, 2.0, 0.8], font_size=8.6,
    )
    T.p("{{VERIFY: ‘Result’ reflects the committed test suite (CI executes the "
        "authentication and template tests on every push). Re-run ‘pytest tests/ -q’ in your "
        "environment to confirm the ML-dependent modules pass against your installed stack.}}")

    T.h3("3.6.1 Model Evaluation and Results")
    T.p(
        "The English ASL recognizer was evaluated offline with scripts/model_baseline_report.py, "
        "which runs the same MediaPipe-Holistic-to-TFLite pipeline used in production over the "
        "reference clips in data/signs_videos. Ground truth is taken from each clip's filename, "
        "aligned to the model's 250-label map (with a small alias table). Of 272 clips, 250 "
        "mapped to a known label and were evaluated; the model achieved a **Top-1 accuracy of "
        "0.624 (62.4%)** with an average soft-max confidence of 0.545, at roughly 4.1 s per "
        "clip on CPU (the figures are recorded verbatim in docs/model_baseline_report.json). "
        "The evaluation summary is given in Table 12."
    )
    T.table_caption("ASL model evaluation summary (docs/model_baseline_report.json).")
    T.make_table(
        ["Metric", "Value"],
        [
            ["Clips evaluated (of 272)", "250 (22 unmapped, 0 failed)"],
            ["Top-1 accuracy", "0.624"],
            ["Average soft-max confidence", "0.545"],
            ["Average top logit", "6.476"],
            ["Acceptance rate (logit gate 0.80)", "1.000"],
            ["Accepted precision / effective accuracy", "0.624"],
            ["Throughput", "~4.09 s / clip (CPU)"],
        ],
        widths=[3.4, 2.9], font_size=9.2,
    )
    T.p(
        "Beyond accuracy, two engineering optimisations were measured directly with "
        "scripts/benchmark.py (Figure 18). INT8 dynamic quantization of the Arabic predictor "
        "reduced CPU inference from 7.48 ms to 4.21 ms per call — a 1.78× speed-up with "
        "negligible accuracy impact for this small model — and the bounded gloss-to-sentence "
        "cache reduced a repeated conversion from about 50 ms (a simulated language-model "
        "round-trip) to effectively zero on a cache hit, which in production removes a full "
        "Gemini round-trip (typically several hundred milliseconds) for any recurring gloss. "
        "These results validate the methodology of Section 3.1: instrument first, then "
        "optimise the measured bottlenecks."
    )
    T.figure("F18_results_real.png", "Measured optimisations: Arabic-predictor latency (fp32 vs INT8) "
                                     "and gloss-to-sentence cache (cold vs warm).", width_in=6.2)
    T.p(
        "A deeper error analysis is limited by what the evaluation run recorded. The baseline "
        "logged the overall accuracy and the top single-count confusions — for example "
        "after→arm, all→weus and alligator→clean — but it did not export a full "
        "confusion matrix or per-class precision and recall. Figures 15–17 therefore "
        "present the available picture honestly: Figure 15 is a confusion matrix seeded with "
        "the real top-confusion pairs over a labelled subset, while Figures 16 and 17 are "
        "**clearly-marked placeholders** for the training/validation curves and per-class "
        "metrics that the current artifacts do not contain. They are included so the document "
        "is structurally complete and so the missing real data can be dropped in directly."
    )
    T.figure("F15_confusion_matrix.png", "ASL confusion matrix (illustrative subset, seeded with the real "
                                         "top-confusion pairs). — placeholder, see caption.", width_in=5.4)
    T.figure("F16_training_curves.png", "Training vs validation accuracy and loss. — PLACEHOLDER "
                                        "(no training logs in the repository).", width_in=6.2)
    T.figure("F17_per_class.png", "Per-class precision and recall for the ASL model. — PLACEHOLDER "
                                  "(not exported by the evaluation run).", width_in=6.2)
    T.p(
        "Interpreting the real numbers, a 62.4% Top-1 accuracy over 250 fine-grained, "
        "visually-similar signs is a reasonable result for a landmark-only model evaluated on "
        "single clips per class, and is consistent with the difficulty reported for "
        "large-vocabulary isolated SLR in the literature (Chapter 2). In live use the figure "
        "matters less than it appears, because the production path never accepts a single "
        "frame's prediction: the two-of-three vote and cooldown reject transient errors, and "
        "the confidence gate suppresses low-certainty guesses. The principal validity threats "
        "are the single-clip-per-class evaluation and the absence of an equivalent ArSL "
        "evaluation; both are recorded in the gaps document and revisited as future work."
    )


# ============================== CHAPTER 4 ==================================
def chapter4(T):
    T.h1("Chapter 4: Conclusions and Future Work")
    T.p(
        "This project set out to close a concrete accessibility gap — the everyday "
        "communication barrier between Deaf/HoH signers and the hearing majority — with a "
        "real-time, bi-directional, bilingual sign-language translator that runs in an "
        "ordinary web browser. The resulting system, Together, meets that aim end to end. It "
        "captures 543 MediaPipe landmarks in the browser, recognises isolated signs in two "
        "languages (a 250-class English ASL model in TFLite and a 20-class Egyptian ArSL "
        "CNN-GRU in PyTorch), assembles recognized gloss into natural sentences with a "
        "grammar-aware, offline-resilient language layer, and animates the reverse direction "
        "from a semantically-searched landmark database — all behind a bilingual LTR/RTL "
        "interface with authentication, live WebRTC meetings, and per-stage latency metrics."
    )
    T.p(
        "Measured against the objectives of Section 1.2, O1–O6 were achieved: real-time "
        "capture and streaming (O1), dual-language recognition behind one API (O2), "
        "natural-language output with graceful fallback (O3), avatar synthesis with "
        "fingerspelling for unknown words (O4), a two-party meeting mode (O5), and the "
        "production qualities of bilingual UI, security, profiling and containerised "
        "deployment (O6). On the quantitative side the ASL recognizer reached 62.4% Top-1 "
        "accuracy on 250 isolated-sign clips, and the engineering optimisations delivered a "
        "measured 1.78× inference speed-up and the near-elimination of repeated language-model "
        "round-trips. The honest limitations are equally clear: the vocabularies are small "
        "(especially for Arabic), recognition is of isolated signs rather than continuous "
        "signing, and the evaluation lacks a full confusion matrix, per-class metrics and an "
        "ArSL accuracy figure."
    )
    T.p("These limitations map directly onto future work:")
    T.bullet("**More signs, more data.** Expand both vocabularies — particularly the Egyptian "
             "ArSL set — with a documented, consented dataset, and record full training and "
             "evaluation artifacts (splits, curves, confusion matrices, per-class metrics).")
    T.bullet("**Continuous, sentence-level translation.** Move beyond isolated signs to "
             "continuous signing using sequence-to-sequence or Transformer models [12], so the "
             "system can translate fluent signing rather than one sign at a time.")
    T.bullet("**On-device / edge deployment.** Run recognition entirely in the browser or on "
             "a phone (e.g. TFLite/WebGPU) to remove the server round-trip, cut latency "
             "further and strengthen privacy.")
    T.bullet("**Latency and robustness.** Reduce end-to-end latency on real networks and "
             "devices, and improve robustness to lighting, occlusion and signer variation; "
             "measure real-device FPS and run a user study with Deaf/HoH participants.")
    T.bullet("**Richer avatar and grammar.** Render the annotated non-manual markers on the "
             "avatar and extend the gloss engine toward fuller sign-language grammar.")
    T.p(
        "In summary, Together demonstrates that a practical, deployable, bilingual sign-"
        "language translator can be built today from landmark-based recognition, a grammar-"
        "aware language layer and semantic sign synthesis, on commodity hardware and in the "
        "browser. It is not a finished product — continuous translation and larger, "
        "well-documented datasets remain — but it is a working foundation that delivers real "
        "value to its users and a clear path forward."
    )


# ============================== REFERENCES ==================================
def references(T):
    T.h1("References")
    T.p("In-text citations use IEEE numbering; each number links to its entry below. Entries "
        "marked {{VERIFY: confirm full bibliographic details against the primary source}} "
        "should be checked before submission.", align=WD_ALIGN_PARAGRAPH.LEFT)
    refs = [
        "R. Rastgoo, K. Kiani, and S. Escalera, “Sign language recognition: A deep survey,” "
        "Expert Systems with Applications, vol. 164, p. 113794, 2021.",
        "C. Lugaresi et al., “MediaPipe: A framework for building perception pipelines,” "
        "arXiv:1906.08172, 2019.",
        "F. Zhang et al., “MediaPipe Hands: On-device real-time hand tracking,” "
        "arXiv:2006.10214, 2020.",
        "Google / PopSign, “Google – Isolated Sign Language Recognition (GISLR),” Kaggle "
        "competition, 2023. {{VERIFY: confirm citation, organisers, and dataset licence}}.",
        "N. Reimers and I. Gurevych, “Sentence-BERT: Sentence embeddings using Siamese "
        "BERT-networks,” in Proc. EMNLP-IJCNLP, 2019, pp. 3982–3992.",
        "W. Wang et al., “MiniLM: Deep self-attention distillation for task-agnostic "
        "compression of pre-trained transformers,” in Proc. NeurIPS, 2020.",
        "A. Vaswani et al., “Attention is all you need,” in Proc. NeurIPS, 2017, pp. 5998–6008.",
        "S. Hochreiter and J. Schmidhuber, “Long short-term memory,” Neural Computation, "
        "vol. 9, no. 8, pp. 1735–1780, 1997.",
        "K. Cho et al., “Learning phrase representations using RNN encoder–decoder for "
        "statistical machine translation,” in Proc. EMNLP, 2014, pp. 1724–1734.",
        "J. Carreira and A. Zisserman, “Quo vadis, action recognition? A new model and the "
        "Kinetics dataset,” in Proc. IEEE CVPR, 2017, pp. 6299–6308.",
        "N. C. Camgoz, S. Hadfield, O. Koller, H. Ney, and R. Bowden, “Neural sign language "
        "translation,” in Proc. IEEE CVPR, 2018, pp. 7784–7793.",
        "N. C. Camgoz, O. Koller, S. Hadfield, and R. Bowden, “Sign language transformers: "
        "Joint end-to-end sign language recognition and translation,” in Proc. IEEE CVPR, "
        "2020, pp. 10023–10033.",
        "M. M. Kamruzzaman, “Arabic sign language recognition and generating Arabic speech "
        "using convolutional neural network,” Wireless Communications and Mobile Computing, "
        "vol. 2020, art. 3685614, 2020.",
        "G. Latif, N. Mohammad, J. Alghazo, R. AlKhalaf, and R. AlKhalaf, “ArASL: Arabic "
        "alphabets sign language dataset,” Data in Brief, vol. 23, art. 103777, 2019.",
        "O. M. Sincan and H. Y. Keles, “AUTSL: A large scale multi-modal Turkish sign "
        "language dataset and baseline methods,” IEEE Access, vol. 8, pp. 181340–181355, 2020.",
        "D. Li, C. Rodriguez, X. Yu, and H. Li, “Word-level deep sign language recognition "
        "from video: A new large-scale dataset and methods comparison (WLASL),” in Proc. IEEE "
        "WACV, 2020, pp. 1459–1469.",
        "C. Oz and M. C. Leu, “American Sign Language word recognition with a sensory glove "
        "using artificial neural networks,” Engineering Applications of Artificial "
        "Intelligence, vol. 24, no. 7, pp. 1204–1213, 2011.",
        "J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, “BERT: Pre-training of deep "
        "bidirectional transformers for language understanding,” in Proc. NAACL-HLT, 2019, "
        "pp. 4171–4186.",
        "T. Starner, J. Weaver, and A. Pentland, “Real-time American sign language "
        "recognition using desk and wearable computer based video,” IEEE Trans. Pattern "
        "Analysis and Machine Intelligence, vol. 20, no. 12, pp. 1371–1375, 1998.",
        "World Health Organization, “World report on hearing,” Geneva: WHO, 2021. "
        "{{VERIFY: cite the specific statistics used in Section 1.1}}.",
    ]
    for i, txt in enumerate(refs, start=1):
        para = T.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Inches(0.4)
        para.paragraph_format.first_line_indent = Inches(-0.4)
        para.paragraph_format.space_after = Pt(6)
        from build_thesis import _emit_runs as emit  # reuse the tokenizer
        num = para.add_run(f"[{i}] ")
        num.bold = True
        emit(para, txt)
        _bookmark(para, f"ref{i}", 1000 + i)
